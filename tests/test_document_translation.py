from docx import Document

from document_translation import translate_docx, translate_markdown


class FakeTranslator:
    def __init__(self):
        self.batches: list[list[str]] = []

    def should_translate(self, text: str) -> bool:
        return bool(text.strip())

    def translate_many_texts(self, texts):
        values = list(texts)
        self.batches.append(values)
        return [f"译:{value}" for value in values]


def test_markdown_preserves_code_links_and_prefixes() -> None:
    translator = FakeTranslator()
    source = "# Heading\n\nUse `x = 1` and [docs](https://example.com).\n```py\nprint('keep')\n```\n"
    result = translate_markdown(translator, source)
    assert result.startswith("# 译:Heading")
    assert "`x = 1`" in result
    assert "[docs](https://example.com)" in result
    assert "print('keep')" in result
    assert len(translator.batches) == 1


def test_docx_preserves_structure_and_translates_visible_text(tmp_path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.add_heading("Heading", level=1)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table text"
    document.sections[0].header.paragraphs[0].text = "Header text"
    document.save(source)

    translate_docx(FakeTranslator(), source, output)
    translated = Document(output)
    assert translated.paragraphs[0].style.name.startswith("Heading")
    assert translated.paragraphs[0].text == "译:Heading"
    assert translated.tables[0].cell(0, 0).text == "译:Table text"
    assert translated.sections[0].header.paragraphs[0].text == "译:Header text"

