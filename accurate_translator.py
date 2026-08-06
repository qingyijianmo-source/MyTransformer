"""Accurate Chinese-English translation for long text and documents.

The original hand-written Transformer remains available in ``run.py``.  This
module uses a pretrained OPUS-MT checkpoint, sentence-aware chunking, batched
beam search, and a durable translation cache for practical document work.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable, Optional, Sequence

import torch
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer


PROJECT_DIR = Path(__file__).resolve().parent
DEFAULT_CONFIG = PROJECT_DIR / "config.accurate.json"
ProgressCallback = Callable[[int, int, str], None]
CJK_PATTERN = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
LATIN_PATTERN = re.compile(r"[A-Za-z]")
MARKDOWN_PREFIX = re.compile(r"^(\s*(?:#{1,6}\s+|[-*+]\s+|\d+[.)]\s+|>\s+))(.*)$")
MARKDOWN_PROTECTED = re.compile(r"(`[^`]*`|https?://\S+|!?\[[^\]]*\]\([^)]*\))")


def normalize_direction(direction: str) -> str:
    aliases = {
        "zh-en": "zh-en",
        "zh_to_en": "zh-en",
        "中译英": "zh-en",
        "en-zh": "en-zh",
        "en_to_zh": "en-zh",
        "英译中": "en-zh",
    }
    normalized = aliases.get(str(direction).strip().lower())
    if normalized is None:
        raise ValueError(f"不支持的翻译方向：{direction}")
    return normalized


def detect_translation_direction(text: str) -> str:
    """Choose a direction from visible script counts; manual selection wins in UI."""

    cjk_count = len(CJK_PATTERN.findall(text))
    latin_count = len(LATIN_PATTERN.findall(text))
    if cjk_count == 0 and latin_count == 0:
        return "zh-en"
    # A Chinese sentence often contains Latin product names, while English
    # prose rarely contains many Han characters.  Weight Han characters more.
    return "zh-en" if cjk_count * 2 >= latin_count else "en-zh"


@dataclass(frozen=True)
class TranslatorSettings:
    direction: str
    model_name: str
    revision: Optional[str]
    fine_tuned_model_dir: Optional[str]
    glossary_file: Optional[str]
    context_rules_file: Optional[str]
    target_prefix: str
    device: str
    batch_size: int
    num_beams: int
    length_penalty: float
    repetition_penalty: float
    no_repeat_ngram_size: int
    max_source_tokens: int
    max_new_tokens: int

    @classmethod
    def from_file(cls, path: Path, direction: str = "zh-en") -> "TranslatorSettings":
        with path.open("r", encoding="utf-8") as file:
            raw = json.load(file)
        if direction not in {"zh-en", "en-zh"}:
            raise ValueError(f"不支持的翻译方向：{direction}")
        direction_values = raw.get("directions", {}).get(direction, {})
        merged = {**raw, **direction_values}
        if "model_name" not in merged:
            raise ValueError(f"配置中缺少 {direction} 的 model_name")
        return cls(
            direction=direction,
            model_name=str(merged["model_name"]),
            revision=merged.get("revision") or None,
            fine_tuned_model_dir=merged.get("fine_tuned_model_dir") or None,
            glossary_file=merged.get("glossary_file") or None,
            context_rules_file=merged.get("context_rules_file") or None,
            target_prefix=str(merged.get("target_prefix", "")),
            device=str(merged.get("device", "auto")),
            batch_size=max(1, int(merged.get("batch_size", 16))),
            num_beams=max(1, int(merged.get("num_beams", 5))),
            length_penalty=float(merged.get("length_penalty", 1.0)),
            repetition_penalty=max(1.0, float(merged.get("repetition_penalty", 1.0))),
            no_repeat_ngram_size=max(0, int(merged.get("no_repeat_ngram_size", 3))),
            max_source_tokens=max(32, int(merged.get("max_source_tokens", 384))),
            max_new_tokens=max(32, int(merged.get("max_new_tokens", 512))),
        )


class TranslationCache:
    """Append-only JSONL cache; an interrupted document run can be resumed."""

    def __init__(self, path: Optional[Path], signature: str):
        self.path = path
        self.signature = signature
        self.values: dict[str, str] = {}
        if path is not None and path.is_file():
            with path.open("r", encoding="utf-8") as file:
                for line in file:
                    try:
                        record = json.loads(line)
                    except (json.JSONDecodeError, UnicodeDecodeError):
                        continue
                    if record.get("signature") == signature:
                        self.values[str(record.get("key", ""))] = str(
                            record.get("translation", "")
                        )

    def _key(self, text: str) -> str:
        return hashlib.sha256(text.encode("utf-8")).hexdigest()

    def get(self, text: str) -> Optional[str]:
        return self.values.get(self._key(text))

    def put(self, text: str, translation: str) -> None:
        key = self._key(text)
        self.values[key] = translation
        if self.path is None:
            return
        self.path.parent.mkdir(parents=True, exist_ok=True)
        record = {
            "signature": self.signature,
            "key": key,
            "translation": translation,
        }
        with self.path.open("a", encoding="utf-8", newline="\n") as file:
            file.write(json.dumps(record, ensure_ascii=False) + "\n")
            file.flush()
            os.fsync(file.fileno())


class AccurateTranslator:
    def __init__(
        self,
        config_path: Path = DEFAULT_CONFIG,
        cache_path: Optional[Path] = None,
        progress: Optional[ProgressCallback] = None,
        direction: str = "zh-en",
    ):
        self.config_path = config_path.resolve()
        self.direction = normalize_direction(direction)
        self.settings = TranslatorSettings.from_file(self.config_path, self.direction)
        self.progress = progress
        self.glossary = self._load_glossary()
        self.context_rules = self._load_context_rules()
        os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")
        self.device = self._select_device(self.settings.device)
        model_source, revision, model_fingerprint = self._resolve_model_source()
        self.active_model = model_source
        self.using_fine_tuned_model = Path(model_source).is_absolute()
        model_label = "本地增强模型" if self.using_fine_tuned_model else model_source
        self._notify(0, 1, f"正在加载 {model_label}…")
        revision_args = {"revision": revision} if revision else {}
        # The fine-tuned model retains the base vocabulary.  Loading its
        # tokenizer from the Hub cache also avoids SentencePiece's unreliable
        # handling of non-ASCII local paths on Windows.
        tokenizer_revision_args = (
            {"revision": self.settings.revision} if self.settings.revision else {}
        )
        self.tokenizer = AutoTokenizer.from_pretrained(
            self.settings.model_name, **tokenizer_revision_args
        )
        dtype = torch.float16 if self.device.type == "cuda" else torch.float32
        self.model = AutoModelForSeq2SeqLM.from_pretrained(
            model_source, torch_dtype=dtype, **revision_args
        ).to(self.device)
        self.model.eval()
        if self.device.type == "cuda":
            torch.backends.cuda.matmul.allow_tf32 = True
        signature_source = json.dumps(
            {
                "pipeline_version": 8,
                "direction": self.direction,
                "model": model_source,
                "model_fingerprint": model_fingerprint,
                "revision": revision,
                "beams": self.settings.num_beams,
                "length_penalty": self.settings.length_penalty,
                "repetition_penalty": self.settings.repetition_penalty,
                "no_repeat_ngram_size": self.settings.no_repeat_ngram_size,
                "glossary": self.glossary,
                "context_rules": self.context_rules,
            },
            sort_keys=True,
        )
        signature = hashlib.sha256(signature_source.encode("utf-8")).hexdigest()
        self.cache = TranslationCache(cache_path, signature)
        self._notify(1, 1, f"{model_label}已加载到 {self.device}")

    @property
    def direction_label(self) -> str:
        return "中译英" if self.direction == "zh-en" else "英译中"

    @property
    def source_label(self) -> str:
        return "中文" if self.direction == "zh-en" else "英文"

    @property
    def target_label(self) -> str:
        return "英文" if self.direction == "zh-en" else "中文"

    def should_translate(self, text: str) -> bool:
        pattern = CJK_PATTERN if self.direction == "zh-en" else LATIN_PATTERN
        return bool(pattern.search(text))

    def _resolve_model_source(self) -> tuple[str, Optional[str], str]:
        """Use a quality-gated local checkpoint when one has been promoted."""

        configured = self.settings.fine_tuned_model_dir
        if configured:
            candidate = Path(configured)
            if not candidate.is_absolute():
                candidate = (self.config_path.parent / candidate).resolve()
            manifest_path = candidate / "training_manifest.json"
            required = (
                candidate / "config.json",
                candidate / "model.safetensors",
                candidate / "source.spm",
                candidate / "target.spm",
                candidate / "vocab.json",
            )
            if manifest_path.is_file() and all(path.is_file() for path in required):
                try:
                    manifest_text = manifest_path.read_text(encoding="utf-8")
                    manifest = json.loads(manifest_text)
                except (OSError, UnicodeDecodeError, json.JSONDecodeError):
                    manifest = {}
                    manifest_text = ""
                if manifest.get("status") == "accepted":
                    fingerprint = hashlib.sha256(
                        manifest_text.encode("utf-8")
                    ).hexdigest()
                    return str(candidate), None, fingerprint
        fallback = f"{self.settings.model_name}@{self.settings.revision or 'default'}"
        return self.settings.model_name, self.settings.revision, fallback

    def _load_glossary(self) -> dict[str, dict[str, object]]:
        configured = self.settings.glossary_file
        if not configured:
            return {}
        path = Path(configured)
        if not path.is_absolute():
            path = (self.config_path.parent / path).resolve()
        if not path.is_file():
            return {}
        with path.open("r", encoding="utf-8") as file:
            values = json.load(file)
        if not isinstance(values, dict):
            raise ValueError(f"术语表必须是 JSON 对象：{path}")
        glossary: dict[str, dict[str, object]] = {}
        for source, raw_rule in values.items():
            source = str(source)
            if isinstance(raw_rule, str):
                target = raw_rule
                aliases: list[str] = []
            elif isinstance(raw_rule, dict):
                target = str(raw_rule.get("target", ""))
                aliases = [
                    str(alias)
                    for alias in raw_rule.get("aliases", [])
                    if str(alias)
                ]
            else:
                raise ValueError(f"术语 {source!r} 的规则必须是字符串或对象")
            if source and target:
                glossary[source] = {"target": target, "aliases": aliases}
        return glossary

    def _load_context_rules(self) -> list[dict[str, object]]:
        """Load source-aware lexical rewrites and guarded output repairs."""

        configured = self.settings.context_rules_file
        if not configured:
            return []
        path = Path(configured)
        if not path.is_absolute():
            path = (self.config_path.parent / path).resolve()
        if not path.is_file():
            return []
        with path.open("r", encoding="utf-8") as file:
            values = json.load(file)
        if not isinstance(values, list):
            raise ValueError(f"上下文消歧规则必须是 JSON 数组：{path}")
        rules: list[dict[str, object]] = []
        for index, raw_rule in enumerate(values, start=1):
            if not isinstance(raw_rule, dict):
                raise ValueError(f"上下文消歧规则 #{index} 必须是对象")
            source = str(raw_rule.get("source", "")).strip()
            target = str(raw_rule.get("target", "")).strip()
            rewrite = str(raw_rule.get("rewrite", source)).strip()
            if not source or not target or not rewrite:
                raise ValueError(f"上下文消歧规则 #{index} 缺少 source、rewrite 或 target")
            rules.append(
                {
                    "source": source,
                    "rewrite": rewrite,
                    "target": target,
                    "aliases": [
                        str(value)
                        for value in raw_rule.get("aliases", [])
                        if str(value)
                    ],
                    "when_any": [
                        str(value).casefold()
                        for value in raw_rule.get("when_any", [])
                        if str(value)
                    ],
                    "unless_any": [
                        str(value).casefold()
                        for value in raw_rule.get("unless_any", [])
                        if str(value)
                    ],
                }
            )
        return rules

    @staticmethod
    def _source_term_pattern(source: str) -> re.Pattern[str]:
        return re.compile(rf"(?<!\w){re.escape(source)}(?!\w)", flags=re.IGNORECASE)

    def _apply_context_rules(self, text: str) -> tuple[str, dict[str, str]]:
        """Disambiguate only when both the source term and its context match."""

        rewritten = text
        original_folded = text.casefold()
        replacements: dict[str, str] = {}
        for rule in self.context_rules:
            source = str(rule["source"])
            pattern = self._source_term_pattern(source)
            if pattern.search(rewritten) is None:
                continue
            when_any = list(rule["when_any"])
            unless_any = list(rule["unless_any"])
            if when_any and not any(value in original_folded for value in when_any):
                continue
            if any(value in original_folded for value in unless_any):
                continue
            rewritten = pattern.sub(str(rule["rewrite"]), rewritten)
            target = str(rule["target"])
            for alias in rule["aliases"]:
                replacements[str(alias)] = target
        return rewritten, replacements

    def _protect_glossary(self, text: str) -> tuple[str, dict[str, str]]:
        protected = text
        replacements: dict[str, str] = {}
        terms = sorted(self.glossary.items(), key=lambda item: len(item[0]), reverse=True)
        # Give the model the approved English term in the source so it can
        # arrange fluent grammar; known alternative outputs are normalized
        # back to the approved term after generation.
        for source, rule in terms:
            source_present = (
                source in protected
                if self.direction == "zh-en"
                else re.search(
                    rf"(?<!\w){re.escape(source)}(?!\w)",
                    protected,
                    flags=re.IGNORECASE,
                )
                is not None
            )
            if not source_present:
                continue
            target = str(rule["target"])
            # Marian zh→en can fluently arrange an approved English term that
            # is injected into Chinese source text.  The reverse direction is
            # less tolerant of Chinese embedded in English and may omit it, so
            # en→zh keeps the source intact and normalizes known outputs only.
            if self.direction == "zh-en":
                protected = protected.replace(source, target)
            for alias in rule["aliases"]:
                replacements[str(alias)] = target
        return protected, replacements

    @staticmethod
    def _restore_glossary(text: str, replacements: dict[str, str]) -> str:
        restored = text
        for alias, target in sorted(
            replacements.items(), key=lambda item: len(item[0]), reverse=True
        ):
            if CJK_PATTERN.search(alias):
                restored = restored.replace(alias, target)
            else:
                restored = re.sub(
                    rf"(?<!\w){re.escape(alias)}(?!\w)",
                    target,
                    restored,
                    flags=re.IGNORECASE,
                )
        return restored

    @staticmethod
    def _select_device(requested: str) -> torch.device:
        if requested == "cuda" and not torch.cuda.is_available():
            raise RuntimeError("配置要求 CUDA，但当前 PyTorch 无法使用 CUDA")
        if requested == "cpu":
            return torch.device("cpu")
        if requested == "cuda" or torch.cuda.is_available():
            return torch.device("cuda")
        return torch.device("cpu")

    def _notify(self, completed: int, total: int, message: str) -> None:
        if self.progress is not None:
            self.progress(completed, max(total, 1), message)

    def _token_count(self, text: str) -> int:
        return len(
            self.tokenizer.encode(
                self.settings.target_prefix + text,
                add_special_tokens=False,
                verbose=False,
            )
        )

    def _sentence_split(self, text: str) -> list[str]:
        if not text:
            return []
        endings = set("。！？!?；;\n")
        if self.direction == "en-zh":
            endings.add(".")
        closers = set("”’」』】）》)]\"")
        pieces: list[str] = []
        start = 0
        index = 0
        while index < len(text):
            if text[index] in endings:
                end = index + 1
                while end < len(text) and text[end] in closers:
                    end += 1
                piece = text[start:end].strip()
                if piece:
                    pieces.append(piece)
                start = end
                index = end
                continue
            index += 1
        remainder = text[start:].strip()
        if remainder:
            pieces.append(remainder)
        return pieces

    def _hard_token_split(self, text: str) -> list[str]:
        token_ids = self.tokenizer.encode(
            text, add_special_tokens=False, verbose=False
        )
        prefix_tokens = self.tokenizer.encode(
            self.settings.target_prefix, add_special_tokens=False, verbose=False
        )
        size = max(1, self.settings.max_source_tokens - len(prefix_tokens))
        return [
            self.tokenizer.decode(token_ids[index : index + size]).strip()
            for index in range(0, len(token_ids), size)
            if token_ids[index : index + size]
        ]

    def split_for_translation(self, text: str) -> list[str]:
        """Keep paragraph context; split only when the model token limit requires it."""

        stripped = text.strip()
        if not stripped:
            return []
        if self._token_count(stripped) <= self.settings.max_source_tokens:
            return [stripped]

        sentence_parts: list[str] = []
        for sentence in self._sentence_split(stripped):
            if self._token_count(sentence) <= self.settings.max_source_tokens:
                sentence_parts.append(sentence)
                continue
            clauses = re.split(r"(?<=[，,、：:])", sentence)
            current = ""
            for clause in clauses:
                candidate = current + clause
                if current and self._token_count(candidate) > self.settings.max_source_tokens:
                    if self._token_count(current) <= self.settings.max_source_tokens:
                        sentence_parts.append(current.strip())
                    else:
                        sentence_parts.extend(self._hard_token_split(current))
                    current = clause
                else:
                    current = candidate
            if current.strip():
                if self._token_count(current) <= self.settings.max_source_tokens:
                    sentence_parts.append(current.strip())
                else:
                    sentence_parts.extend(self._hard_token_split(current))

        separator = " " if self.direction == "en-zh" else ""
        result: list[str] = []
        current = ""
        for part in sentence_parts:
            candidate = part if not current else current + separator + part
            if current and self._token_count(candidate) > self.settings.max_source_tokens:
                result.append(current)
                current = part
            else:
                current = candidate
        if current:
            result.append(current)
        return [piece for piece in result if piece]

    def _generate_batch(self, texts: Sequence[str]) -> list[str]:
        model_inputs = [self.settings.target_prefix + text for text in texts]
        encoded = self.tokenizer(
            model_inputs,
            return_tensors="pt",
            padding=True,
            truncation=True,
            max_length=self.settings.max_source_tokens + 8,
        ).to(self.device)
        generation_options = {
            "num_beams": self.settings.num_beams,
            "max_new_tokens": self.settings.max_new_tokens,
            "length_penalty": self.settings.length_penalty,
            "repetition_penalty": self.settings.repetition_penalty,
            "early_stopping": True,
            "no_repeat_ngram_size": self.settings.no_repeat_ngram_size,
            "renormalize_logits": True,
        }
        with torch.inference_mode():
            generated = self.model.generate(**encoded, **generation_options)
        return [
            value.strip()
            for value in self.tokenizer.batch_decode(
                generated, skip_special_tokens=True, clean_up_tokenization_spaces=True
            )
        ]

    def translate_segments(self, segments: Sequence[str]) -> list[str]:
        if not segments:
            return []
        translated: dict[str, str] = {}
        missing: list[str] = []
        seen: set[str] = set()
        for segment in segments:
            cached = self.cache.get(segment)
            if cached is not None:
                translated[segment] = cached
            elif segment not in seen:
                seen.add(segment)
                missing.append(segment)

        total = len(missing)
        cached_count = len(segments) - sum(1 for segment in segments if segment in missing)
        if total == 0:
            self._notify(len(segments), len(segments), "全部内容已从断点缓存恢复")
            return [translated[segment] for segment in segments]

        batch_size = self.settings.batch_size
        for start in range(0, total, batch_size):
            batch = missing[start : start + batch_size]
            outputs = self._generate_batch(batch)
            for source, output in zip(batch, outputs):
                translated[source] = output
                self.cache.put(source, output)
            completed = min(start + len(batch), total)
            self._notify(
                completed,
                total,
                f"正在翻译 {completed}/{total} 个片段（另有 {cached_count} 个缓存命中）",
            )
        return [translated[segment] for segment in segments]

    def translate_many_texts(self, texts: Sequence[str]) -> list[str]:
        plans: list[tuple[str, str, list[str], str, dict[str, str]]] = []
        all_segments: list[str] = []
        for text in texts:
            if not text or not text.strip():
                plans.append(("", text, [], "", {}))
                continue
            leading_match = re.match(r"^\s*", text)
            trailing_match = re.search(r"\s*$", text)
            leading = leading_match.group(0) if leading_match else ""
            trailing = trailing_match.group(0) if trailing_match else ""
            end = len(text) - len(trailing) if trailing else len(text)
            core = text[len(leading) : end]
            if not self.should_translate(core):
                segments: list[str] = []
                replacements: dict[str, str] = {}
            else:
                contextual_core, context_replacements = self._apply_context_rules(core)
                protected_core, glossary_replacements = self._protect_glossary(
                    contextual_core
                )
                replacements = {**glossary_replacements, **context_replacements}
                segments = self.split_for_translation(protected_core)
                all_segments.extend(segments)
            plans.append((leading, core, segments, trailing, replacements))

        outputs = iter(self.translate_segments(all_segments))
        results: list[str] = []
        for leading, core, segments, trailing, replacements in plans:
            if not segments:
                results.append(leading + core + trailing)
                continue
            translated_segments = [next(outputs) for _ in segments]
            separator = " " if self.direction == "zh-en" else ""
            translated_core = self._normalize_translation(
                self._restore_glossary(
                    separator.join(translated_segments), replacements
                )
            )
            results.append(leading + translated_core + trailing)
        return results

    def _normalize_translation(self, text: str) -> str:
        normalized = text.strip()
        if self.direction != "en-zh":
            return normalized
        # Marian sometimes leaves spaces between Chinese characters or around
        # Chinese punctuation.  Removing only CJK-to-CJK spaces keeps Latin
        # product names and numbers readable.
        normalized = re.sub(
            rf"(?<=[{CJK_PATTERN.pattern[1:-1]}])\s+(?=[{CJK_PATTERN.pattern[1:-1]}])",
            "",
            normalized,
        )
        normalized = re.sub(r"\s+([，。！？；：、])", r"\1", normalized)
        normalized = re.sub(r"([，。！？；：、])\s+", r"\1", normalized)
        normalized = re.sub(r"(?<!\d),(?!\d)", "，", normalized)
        normalized = normalized.replace(";", "；")
        normalized = re.sub(
            r"一块(?=[^，。！？；]{0,16}(?:帷幕|面纱))",
            "一层",
            normalized,
        )
        # Collapse only three-or-more adjacent copies, avoiding legitimate
        # two-character forms such as 人人 or 常常.
        normalized = re.sub(
            r"(?P<phrase>[\u3400-\u4dbf\u4e00-\u9fff]{2,6})(?:的?\s*(?P=phrase)){2,}",
            r"\g<phrase>",
            normalized,
        )
        return normalized

    def translate_text(self, text: str) -> str:
        lines = text.splitlines(keepends=True)
        if not lines:
            return ""
        bodies: list[str] = []
        endings: list[str] = []
        for line in lines:
            body = line.rstrip("\r\n")
            bodies.append(body)
            endings.append(line[len(body) :])
        translated = self.translate_many_texts(bodies)
        return "".join(value + ending for value, ending in zip(translated, endings))


def default_output_path(source: Path, direction: str = "zh-en") -> Path:
    language_suffix = ".en" if normalize_direction(direction) == "zh-en" else ".zh"
    return source.with_name(f"{source.stem}{language_suffix}{source.suffix}")


def default_cache_path(output: Path) -> Path:
    return output.with_name(f"{output.name}.translation-cache.jsonl")


def _atomic_write_text(output: Path, text: str) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(f".{output.name}.tmp")
    temporary.write_text(text, encoding="utf-8")
    os.replace(temporary, output)


def _iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_docx_paragraphs(document) -> Iterable:
    # Keep the OOXML element objects themselves alive.  Storing only id(...)
    # allows CPython to reuse an id after a temporary paragraph wrapper is
    # collected, which can incorrectly suppress an unrelated later paragraph.
    seen: set[object] = set()

    def emit(paragraphs: Iterable):
        for paragraph in paragraphs:
            identity = paragraph._p
            if identity not in seen:
                seen.add(identity)
                yield paragraph

    yield from emit(document.paragraphs)
    for table in document.tables:
        yield from emit(_iter_table_paragraphs(table))
    for section in document.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            yield from emit(container.paragraphs)
            for table in container.tables:
                yield from emit(_iter_table_paragraphs(table))


def _docx_paragraph_text(paragraph) -> str:
    return "".join(node.text or "" for node in paragraph._p.xpath(".//w:t"))


def _replace_docx_paragraph_text(paragraph, translation: str) -> None:
    text_nodes = paragraph._p.xpath(".//w:t")
    if text_nodes:
        text_nodes[0].text = translation
        for node in text_nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(translation)


def document_text_sample(source: Path, max_characters: int = 50_000) -> str:
    """Extract enough visible text to auto-detect a document's direction."""

    suffix = source.suffix.lower()
    if suffix == ".docx":
        from docx import Document

        document = Document(source)
        values: list[str] = []
        length = 0
        for paragraph in _iter_docx_paragraphs(document):
            text = _docx_paragraph_text(paragraph)
            if not text:
                continue
            values.append(text)
            length += len(text)
            if length >= max_characters:
                break
        return "\n".join(values)[:max_characters]
    if suffix in {".txt", ".md", ".markdown"}:
        with source.open("r", encoding="utf-8-sig") as file:
            return file.read(max_characters)
    raise ValueError("目前支持 .txt、.md、.markdown 和 .docx")


def detect_document_direction(source: Path) -> str:
    return detect_translation_direction(document_text_sample(source))


def translate_docx(
    translator: AccurateTranslator,
    source: Path,
    output: Path,
) -> None:
    from docx import Document

    document = Document(source)
    targets = []
    source_texts = []
    for paragraph in _iter_docx_paragraphs(document):
        # Leave Word fields (TOC, page numbers, cross-references) intact.
        if paragraph._p.xpath(".//w:instrText"):
            continue
        text = _docx_paragraph_text(paragraph)
        if translator.should_translate(text):
            targets.append(paragraph)
            source_texts.append(text)
    translated = translator.translate_many_texts(source_texts)
    for paragraph, value in zip(targets, translated):
        _replace_docx_paragraph_text(paragraph, value)
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(f".{output.stem}.tmp{output.suffix}")
    document.save(temporary)
    os.replace(temporary, output)


def translate_markdown(
    translator: AccurateTranslator,
    source_text: str,
) -> str:
    lines = source_text.splitlines(keepends=True)
    output_lines: list[str] = []
    in_fence = False
    for line in lines:
        body = line.rstrip("\r\n")
        ending = line[len(body) :]
        if body.lstrip().startswith("```"):
            in_fence = not in_fence
            output_lines.append(line)
            continue
        if in_fence or not translator.should_translate(body):
            output_lines.append(line)
            continue
        prefix_match = MARKDOWN_PREFIX.match(body)
        prefix, content = (
            (prefix_match.group(1), prefix_match.group(2))
            if prefix_match
            else ("", body)
        )
        pieces = MARKDOWN_PROTECTED.split(content)
        candidates = [
            piece
            for index, piece in enumerate(pieces)
            if index % 2 == 0 and translator.should_translate(piece)
        ]
        translations = iter(translator.translate_many_texts(candidates))
        rebuilt = []
        for index, piece in enumerate(pieces):
            if index % 2 == 0 and translator.should_translate(piece):
                rebuilt.append(next(translations))
            else:
                rebuilt.append(piece)
        output_lines.append(prefix + "".join(rebuilt) + ending)
    return "".join(output_lines)


def translate_document(
    source: Path,
    output: Optional[Path] = None,
    config_path: Path = DEFAULT_CONFIG,
    overwrite: bool = False,
    progress: Optional[ProgressCallback] = None,
    translator: Optional[AccurateTranslator] = None,
    direction: str = "auto",
) -> Path:
    source = source.resolve()
    if not source.is_file():
        raise FileNotFoundError(f"找不到输入文件：{source}")
    if translator is not None:
        resolved_direction = translator.direction
    elif direction == "auto":
        resolved_direction = detect_document_direction(source)
    else:
        resolved_direction = normalize_direction(direction)
    output = (output or default_output_path(source, resolved_direction)).resolve()
    if source == output:
        raise ValueError("输出文件不能覆盖输入文件")
    if output.exists() and not overwrite:
        raise FileExistsError(f"输出已存在：{output}；使用 --overwrite 才能替换")
    suffix = source.suffix.lower()
    if suffix not in {".txt", ".md", ".markdown", ".docx"}:
        raise ValueError("目前支持 .txt、.md、.markdown 和 .docx")
    cache_path = default_cache_path(output)
    if translator is None:
        translator = AccurateTranslator(
            config_path, cache_path, progress, direction=resolved_direction
        )
    else:
        translator.progress = progress
        translator.cache = TranslationCache(cache_path, translator.cache.signature)
    if suffix == ".docx":
        translate_docx(translator, source, output)
    else:
        source_text = source.read_text(encoding="utf-8-sig")
        translated = (
            translate_markdown(translator, source_text)
            if suffix in {".md", ".markdown"}
            else translator.translate_text(source_text)
        )
        _atomic_write_text(output, translated)
    if progress is not None:
        progress(1, 1, f"翻译完成：{output}")
    return output


def _console_progress(completed: int, total: int, message: str) -> None:
    percent = 100.0 * completed / max(total, 1)
    print(f"[{percent:6.2f}%] {message}", flush=True)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--config", type=Path, default=DEFAULT_CONFIG)
    parser.add_argument(
        "--direction",
        choices=("auto", "zh-en", "en-zh"),
        default="auto",
        help="翻译方向；auto 按输入文字识别",
    )
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--text", help="翻译一段文字并退出")
    group.add_argument("--input", type=Path, help="输入 TXT、Markdown 或 DOCX")
    parser.add_argument("--output", type=Path, help="文档输出路径")
    parser.add_argument("--overwrite", action="store_true")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.input is not None:
        output = translate_document(
            args.input,
            args.output,
            args.config,
            args.overwrite,
            _console_progress,
            direction=args.direction,
        )
        print(output)
        return
    if args.text is not None:
        direction = (
            detect_translation_direction(args.text)
            if args.direction == "auto"
            else args.direction
        )
        translator = AccurateTranslator(
            args.config, progress=_console_progress, direction=direction
        )
        print(translator.translate_text(args.text))
        return

    print("===== 高精度中英文互译：可粘贴多行文字 =====")
    print("单独输入 END 开始翻译；输入 quit 退出。")
    translator: Optional[AccurateTranslator] = None
    while True:
        lines: list[str] = []
        print("原文 >>>", flush=True)
        while True:
            try:
                line = input()
            except (EOFError, KeyboardInterrupt):
                print("\nBye!")
                return
            if line.strip().lower() == "quit" and not lines:
                print("Bye!")
                return
            if line.strip() == "END":
                break
            lines.append(line)
        text = "\n".join(lines)
        if text.strip():
            direction = (
                detect_translation_direction(text)
                if args.direction == "auto"
                else args.direction
            )
            if translator is None or translator.direction != direction:
                if translator is not None:
                    translator.model.to("cpu")
                    del translator
                    if torch.cuda.is_available():
                        torch.cuda.empty_cache()
                translator = AccurateTranslator(
                    args.config, progress=_console_progress, direction=direction
                )
            print(f"{translator.target_label} >>>")
            print(translator.translate_text(text))


if __name__ == "__main__":
    try:
        main()
    except (FileNotFoundError, FileExistsError, ValueError, RuntimeError) as error:
        raise SystemExit(f"Error: {error}") from error
