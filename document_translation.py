"""Format-preserving TXT, Markdown, and DOCX translation helpers."""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Iterable, Protocol, Sequence

from translation_quality import detect_translation_direction


MARKDOWN_PREFIX = re.compile(r"^(\s*(?:#{1,6}\s+|[-*+]\s+|\d+[.)]\s+|>\s+))(.*)$")
MARKDOWN_PROTECTED = re.compile(r"(`[^`]*`|https?://\S+|!?\[[^\]]*\]\([^)]*\))")


class DocumentTranslator(Protocol):
    def should_translate(self, text: str) -> bool: ...
    def translate_many_texts(self, texts: Sequence[str]) -> list[str]: ...


def default_output_path(source: Path, direction: str = "zh-en") -> Path:
    normalized = str(direction).strip().lower()
    language_suffix = ".en" if normalized in {"zh-en", "zh_to_en", "中译英"} else ".zh"
    return source.with_name(f"{source.stem}{language_suffix}{source.suffix}")


def default_cache_path(output: Path) -> Path:
    return output.with_name(f"{output.name}.translation-cache.jsonl")


def atomic_write_text(output: Path, text: str) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(f".{output.name}.tmp")
    temporary.write_text(text, encoding="utf-8")
    os.replace(temporary, output)


def _iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def iter_docx_paragraphs(document) -> Iterable:
    seen: set[object] = set()

    def emit(paragraphs: Iterable):
        for paragraph in paragraphs:
            identity = paragraph._p
            if identity not in seen:
                seen.add(identity)
                yield paragraph

    yield from emit(document.paragraphs)
    for table in document.tables:
        yield from emit(_iter_table_paragraphs(table))
    for section in document.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            yield from emit(container.paragraphs)
            for table in container.tables:
                yield from emit(_iter_table_paragraphs(table))


def docx_paragraph_text(paragraph) -> str:
    return "".join(node.text or "" for node in paragraph._p.xpath(".//w:t"))


def replace_docx_paragraph_text(paragraph, translation: str) -> None:
    text_nodes = paragraph._p.xpath(".//w:t")
    if text_nodes:
        text_nodes[0].text = translation
        for node in text_nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(translation)


def document_text_sample(source: Path, max_characters: int = 50_000) -> str:
    suffix = source.suffix.lower()
    if suffix == ".docx":
        from docx import Document

        document = Document(source)
        values: list[str] = []
        length = 0
        for paragraph in iter_docx_paragraphs(document):
            text = docx_paragraph_text(paragraph)
            if not text:
                continue
            values.append(text)
            length += len(text)
            if length >= max_characters:
                break
        return "\n".join(values)[:max_characters]
    if suffix in {".txt", ".md", ".markdown"}:
        with source.open("r", encoding="utf-8-sig") as file:
            return file.read(max_characters)
    raise ValueError("目前支持 .txt、.md、.markdown 和 .docx")


def detect_document_direction(source: Path) -> str:
    return detect_translation_direction(document_text_sample(source))


def translate_docx(translator: DocumentTranslator, source: Path, output: Path) -> None:
    from docx import Document

    document = Document(source)
    targets = []
    source_texts = []
    for paragraph in iter_docx_paragraphs(document):
        if paragraph._p.xpath(".//w:instrText"):
            continue
        text = docx_paragraph_text(paragraph)
        if translator.should_translate(text):
            targets.append(paragraph)
            source_texts.append(text)
    translated = translator.translate_many_texts(source_texts)
    for paragraph, value in zip(targets, translated):
        replace_docx_paragraph_text(paragraph, value)
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(f".{output.stem}.tmp{output.suffix}")
    document.save(temporary)
    os.replace(temporary, output)


def translate_markdown(translator: DocumentTranslator, source_text: str) -> str:
    lines = source_text.splitlines(keepends=True)
    plans: list[tuple[str, list[str], str, set[int]] | str] = []
    candidates: list[str] = []
    in_fence = False
    for line in lines:
        body = line.rstrip("\r\n")
        ending = line[len(body) :]
        if body.lstrip().startswith("```"):
            in_fence = not in_fence
            plans.append(line)
            continue
        if in_fence or not translator.should_translate(body):
            plans.append(line)
            continue
        prefix_match = MARKDOWN_PREFIX.match(body)
        prefix, content = (
            (prefix_match.group(1), prefix_match.group(2))
            if prefix_match
            else ("", body)
        )
        pieces = MARKDOWN_PROTECTED.split(content)
        translated_indices: set[int] = set()
        for index, piece in enumerate(pieces):
            if index % 2 == 0 and translator.should_translate(piece):
                translated_indices.add(index)
                candidates.append(piece)
        plans.append((prefix, pieces, ending, translated_indices))

    translations = iter(translator.translate_many_texts(candidates))
    output_lines: list[str] = []
    for plan in plans:
        if isinstance(plan, str):
            output_lines.append(plan)
            continue
        prefix, pieces, ending, translated_indices = plan
        rebuilt = [
            next(translations) if index in translated_indices else piece
            for index, piece in enumerate(pieces)
        ]
        output_lines.append(prefix + "".join(rebuilt) + ending)
    return "".join(output_lines)
